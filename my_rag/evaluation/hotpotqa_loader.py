"""
HotpotQA 数据集加载器

面试考点：
- 使用真实公开数据集（而非 LLM 造数据）做评估，结果更客观可信
- HotpotQA 是多跳推理数据集：一个问题需要跨多篇文章推理才能回答
- context 字段包含 10 篇文章（2 篇支撑 + 8 篇干扰），模拟真实检索场景
- 将 context 生成 docx 文档，方便后续调试 chunk 策略

数据集字段说明：
- id: 唯一标识
- question: 问题
- answer: 标准答案
- type: "bridge"（桥接推理）或 "comparison"（比较推理）
- level: "easy" / "medium" / "hard"
- context: [{title: str, sentences: [str]}] - 10 篇文章（含干扰项）
- supporting_facts: [{title: str, sent_id: int}] - 支撑句子的位置标注
"""

import os
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")

from my_rag.evaluation.dataset import EvalDataset, EvalSample
from my_rag.utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class HotpotQASample:
    """单条 HotpotQA 样本（解析后的结构化数据）"""
    id: str
    question: str
    answer: str
    type: str                          # "bridge" | "comparison"
    level: str                         # "easy" | "medium" | "hard"
    context_titles: list[str] = field(default_factory=list)
    context_texts: list[str] = field(default_factory=list)   # 每篇文章的完整文本
    supporting_titles: list[str] = field(default_factory=list)  # 支撑文章标题


@dataclass
class HotpotQALoaderConfig:
    """加载配置"""
    split: str = "train"
    num_samples: int = 20              # 取多少条样本
    config_name: str = "distractor"   # "distractor" 或 "fullwiki"
    only_supporting: bool = False      # True：只保留支撑文章；False：保留全部 10 篇（含干扰）
    output_dir: Optional[str] = None  # docx 输出目录，None 则不生成
    save_eval_json: bool = True        # 是否保存评估数据集 JSON


class HotpotQALoader:
    """
    HotpotQA 数据集加载器

    职责：
    1. 从 HuggingFace 拉取指定数量的 HotpotQA 样本
    2. 解析 context/question/answer 字段
    3. 将 context 文章写入 docx 文档（每篇文章一个文件，方便 chunk 调试）
    4. 生成标准 EvalDataset（question + ground_truth + knowledge_base_id）
    """

    def __init__(self, config: HotpotQALoaderConfig | None = None):
        self._config = config or HotpotQALoaderConfig()

    def load(self) -> tuple[list[HotpotQASample], EvalDataset]:
        """
        加载并解析 HotpotQA 数据集

        Returns:
            (samples, eval_dataset)
            - samples: 解析后的结构化样本列表（含 context 文本）
            - eval_dataset: 可直接用于评估的 EvalDataset
        """
        from datasets import load_dataset

        split_str = f"{self._config.split}[:{self._config.num_samples}]"
        logger.info(
            "loading_hotpotqa",
            config=self._config.config_name,
            split=split_str,
            num_samples=self._config.num_samples,
        )

        ds = load_dataset(
            "hotpot_qa",
            self._config.config_name,
            split=split_str,
            trust_remote_code=True,
        )

        samples = []
        for row in ds:
            sample = self._parse_row(row)
            samples.append(sample)

        logger.info("hotpotqa_loaded", total=len(samples))

        eval_dataset = self._build_eval_dataset(samples)

        if self._config.output_dir:
            self._export_docx(samples, Path(self._config.output_dir))

        return samples, eval_dataset

    # ── 内部方法 ──────────────────────────────────────────────────────

    def _parse_row(self, row: dict) -> HotpotQASample:
        """解析单行数据"""
        # context 字段结构：{"title": [...], "sentences": [[str, ...], ...]}
        ctx = row["context"]
        titles: list[str] = ctx["title"]
        sentences_list: list[list[str]] = ctx["sentences"]

        # 每篇文章：将句子列表拼接为完整段落
        context_texts = [
            " ".join(sents).strip()
            for sents in sentences_list
        ]

        # 支撑文章标题（去重）
        sf = row.get("supporting_facts", {})
        supporting_titles = list(dict.fromkeys(sf.get("title", [])))

        if self._config.only_supporting:
            # 只保留支撑文章
            filtered_titles, filtered_texts = [], []
            for title, text in zip(titles, context_texts):
                if title in supporting_titles:
                    filtered_titles.append(title)
                    filtered_texts.append(text)
            titles, context_texts = filtered_titles, filtered_texts

        return HotpotQASample(
            id=row["id"],
            question=row["question"],
            answer=row["answer"],
            type=row.get("type", ""),
            level=row.get("level", ""),
            context_titles=titles,
            context_texts=context_texts,
            supporting_titles=supporting_titles,
        )

    def _build_eval_dataset(self, samples: list[HotpotQASample]) -> EvalDataset:
        """构建 EvalDataset（question + ground_truth，knowledge_base_id 留空待填充）"""
        dataset = EvalDataset(name=f"hotpotqa_{self._config.config_name}_{self._config.num_samples}")
        for s in samples:
            dataset.samples.append(EvalSample(
                question=s.question,
                ground_truth=s.answer,
                knowledge_base_id="",   # 上传文档后由调用方填充
                metadata={
                    "hotpotqa_id": s.id,
                    "type": s.type,
                    "level": s.level,
                    "supporting_titles": s.supporting_titles,
                },
            ))
        return dataset

    def _export_docx(self, samples: list[HotpotQASample], output_dir: Path) -> list[Path]:
        """
        将 context 文章导出为 docx 文件

        策略：将所有样本的 context 文章按 title 去重合并，
        每篇文章生成一个 docx 文件，文件名为 <title>.docx。
        这样上传到知识库后，chunk 调试更直观。

        Returns:
            生成的 docx 文件路径列表
        """
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
        except ImportError:
            logger.error("python-docx not installed, run: pip install python-docx")
            return []

        output_dir.mkdir(parents=True, exist_ok=True)

        # 按 title 去重，合并所有 context 文章
        article_map: dict[str, str] = {}
        for sample in samples:
            for title, text in zip(sample.context_titles, sample.context_texts):
                if title not in article_map:
                    article_map[title] = text

        generated: list[Path] = []
        for title, text in article_map.items():
            safe_name = _safe_filename(title)
            docx_path = output_dir / f"{safe_name}.docx"

            doc = DocxDocument()
            # 标题
            heading = doc.add_heading(title, level=1)
            heading.runs[0].font.size = Pt(16)
            # 正文
            doc.add_paragraph(text)

            doc.save(str(docx_path))
            generated.append(docx_path)

        logger.info(
            "docx_exported",
            total_articles=len(article_map),
            output_dir=str(output_dir),
        )
        return generated

    def export_docx_for_kb(
        self,
        samples: list[HotpotQASample],
        output_dir: Path,
        merge_into_one: bool = False,
    ) -> list[Path]:
        """
        公开方法：导出 docx 文件

        Args:
            samples: 样本列表
            output_dir: 输出目录
            merge_into_one: True 则将所有文章合并为一个 docx；
                            False 则每篇文章一个 docx（推荐，方便 chunk 调试）
        """
        if merge_into_one:
            return self._export_merged_docx(samples, output_dir)
        return self._export_docx(samples, output_dir)

    def _export_merged_docx(self, samples: list[HotpotQASample], output_dir: Path) -> list[Path]:
        """将所有文章合并为一个 docx"""
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
        except ImportError:
            logger.error("python-docx not installed")
            return []

        output_dir.mkdir(parents=True, exist_ok=True)

        article_map: dict[str, str] = {}
        for sample in samples:
            for title, text in zip(sample.context_titles, sample.context_texts):
                if title not in article_map:
                    article_map[title] = text

        doc = DocxDocument()
        doc.add_heading("HotpotQA Context Articles", level=0)

        for title, text in article_map.items():
            doc.add_heading(title, level=1)
            doc.add_paragraph(text)
            doc.add_paragraph("")  # 空行分隔

        merged_path = output_dir / "hotpotqa_contexts_merged.docx"
        doc.save(str(merged_path))
        logger.info("merged_docx_exported", path=str(merged_path), articles=len(article_map))
        return [merged_path]


def _safe_filename(name: str, max_len: int = 80) -> str:
    """将文章标题转为合法文件名"""
    import re
    safe = re.sub(r'[\\/:*?"<>|]', "_", name)
    safe = safe.strip(". ")
    return safe[:max_len] if len(safe) > max_len else safe


def load_hotpotqa_as_eval_dataset(
    num_samples: int = 20,
    split: str = "train",
    config_name: str = "distractor",
    only_supporting: bool = False,
    output_docx_dir: Optional[str] = None,
) -> tuple[list[HotpotQASample], EvalDataset]:
    """
    便捷函数：一行代码加载 HotpotQA 评估数据集

    Args:
        num_samples: 取多少条样本
        split: "train" / "validation"
        config_name: "distractor" / "fullwiki"
        only_supporting: 是否只保留支撑文章（True 更干净，False 含干扰项更真实）
        output_docx_dir: 若指定，则将 context 文章导出为 docx 到此目录

    Returns:
        (samples, eval_dataset)
    """
    cfg = HotpotQALoaderConfig(
        split=split,
        num_samples=num_samples,
        config_name=config_name,
        only_supporting=only_supporting,
        output_dir=output_docx_dir,
    )
    loader = HotpotQALoader(cfg)
    return loader.load()
