"""
Word 文档解析器

面试考点：
- python-docx 库的使用
- 段落 + 表格混合提取
"""

from my_rag.domain.models import ParsedDocument
from my_rag.domain.parser.base import BaseParser


class DocxParser(BaseParser):

    def parse(self, file_path: str) -> ParsedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables_text: list[str] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                tables_text.append("\n".join(rows))

        parts = paragraphs
        if tables_text:
            parts.append("\n--- 表格内容 ---\n")
            parts.extend(tables_text)

        content = "\n\n".join(parts)
        return ParsedDocument(
            content=content,
            metadata={
                "source": file_path,
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "format": "docx",
            },
        )

    def supported_extensions(self) -> list[str]:
        return [".docx"]
